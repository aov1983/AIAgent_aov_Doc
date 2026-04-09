"""
Модуль парсинга документов.

Поддерживаемые форматы:
- MS Word (.docx)
- Текст (.txt)
- Изображения (с OCR)
- Таблицы
"""

from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
import re


@dataclass
class DocumentElement:
    """Элемент документа."""
    element_type: str  # 'heading', 'paragraph', 'table', 'image', 'list_item'
    content: str
    level: int = 0  # Уровень заголовка
    metadata: Dict[str, Any] = None
    page_number: Optional[int] = None
    position: Optional[int] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class ParsedDocument:
    """Результат парсинга документа."""
    title: str
    elements: List[DocumentElement]
    metadata: Dict[str, Any] = None
    source_path: Optional[str] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentParser:
    """Парсер документов."""
    
    def __init__(self):
        self._docx_parser = None
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        Парсинг документа из файла.
        
        Args:
            file_path: Путь к файлу документа
            
        Returns:
            ParsedDocument: Распарсенный документ
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix == '.docx':
            return self._parse_docx(path)
        elif suffix == '.txt':
            return self._parse_txt(path)
        elif suffix in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
            return self._parse_image(path)
        else:
            # Попытка прочитать как текст
            return self._parse_txt(path)
    
    def _parse_docx(self, path: Path) -> ParsedDocument:
        """Парсинг DOCX файла."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Установите python-docx: pip install python-docx")
        
        doc = Document(str(path))
        elements = []
        title = ""
        
        # Извлечение заголовка
        if doc.paragraphs:
            title = doc.paragraphs[0].text.strip() or "Без названия"
        
        position = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style = para.style.name if para.style else ""
            
            # Определение типа элемента
            if style.startswith('Heading'):
                level = int(style.replace('Heading', '').strip()) if style[7:].strip().isdigit() else 1
                element = DocumentElement(
                    element_type='heading',
                    content=text,
                    level=level,
                    position=position
                )
            elif para._element.xpath('.//w:pPr//w:numPr'):
                # Список
                element = DocumentElement(
                    element_type='list_item',
                    content=text,
                    position=position
                )
            else:
                # Обычный абзац
                element = DocumentElement(
                    element_type='paragraph',
                    content=text,
                    position=position
                )
            
            elements.append(element)
            position += 1
        
        # Извлечение таблиц
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            table_element = DocumentElement(
                element_type='table',
                content=self._table_to_text(table_data),
                metadata={'rows': len(table_data), 'columns': len(table_data[0]) if table_data else 0},
                position=position
            )
            elements.append(table_element)
            position += 1
        
        return ParsedDocument(
            title=title,
            elements=elements,
            metadata={'source': str(path), 'type': 'docx'},
            source_path=str(path)
        )
    
    def _parse_txt(self, path: Path) -> ParsedDocument:
        """Парсинг текстового файла."""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        elements = []
        lines = content.split('\n')
        title = lines[0].strip() if lines else "Без названия"
        
        position = 0
        current_paragraph = []
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                if current_paragraph:
                    text = ' '.join(current_paragraph)
                    element = DocumentElement(
                        element_type='paragraph',
                        content=text,
                        position=position
                    )
                    elements.append(element)
                    current_paragraph = []
                    position += 1
                continue
            
            # Проверка на заголовок (Markdown стиль)
            if stripped.startswith('#'):
                if current_paragraph:
                    text = ' '.join(current_paragraph)
                    element = DocumentElement(
                        element_type='paragraph',
                        content=text,
                        position=position
                    )
                    elements.append(element)
                    current_paragraph = []
                    position += 1
                
                level = len(re.match(r'^#+', stripped).group())
                content_text = stripped[level:].strip()
                element = DocumentElement(
                    element_type='heading',
                    content=content_text,
                    level=min(level, 6),
                    position=position
                )
                elements.append(element)
                position += 1
            else:
                current_paragraph.append(stripped)
        
        # Добавление последнего абзаца
        if current_paragraph:
            text = ' '.join(current_paragraph)
            element = DocumentElement(
                element_type='paragraph',
                content=text,
                position=position
            )
            elements.append(element)
        
        return ParsedDocument(
            title=title,
            elements=elements,
            metadata={'source': str(path), 'type': 'txt'},
            source_path=str(path)
        )
    
    def _parse_image(self, path: Path) -> ParsedDocument:
        """Парсинг изображения с OCR."""
        try:
            from PIL import Image
            import pytesseract
        except ImportError:
            raise ImportError("Установите Pillow и pytesseract: pip install Pillow pytesseract")
        
        image = Image.open(path)
        text = pytesseract.image_to_string(image, lang='rus+eng')
        
        elements = []
        if text.strip():
            element = DocumentElement(
                element_type='paragraph',
                content=text,
                metadata={'source_type': 'image_ocr'},
                position=0
            )
            elements.append(element)
        
        return ParsedDocument(
            title=path.stem,
            elements=elements,
            metadata={'source': str(path), 'type': 'image'},
            source_path=str(path)
        )
    
    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """Преобразование таблицы в текстовое представление."""
        if not table_data:
            return ""
        
        lines = []
        for row in table_data:
            lines.append(" | ".join(row))
        return "\n".join(lines)
    
    def parse_content(self, content: str, title: str = "Документ") -> ParsedDocument:
        """
        Парсинг строки контента.
        
        Args:
            content: Текстовое содержимое
            title: Заголовок документа
            
        Returns:
            ParsedDocument: Распарсенный документ
        """
        elements = []
        lines = content.split('\n')
        
        position = 0
        current_paragraph = []
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                if current_paragraph:
                    text = ' '.join(current_paragraph)
                    element = DocumentElement(
                        element_type='paragraph',
                        content=text,
                        position=position
                    )
                    elements.append(element)
                    current_paragraph = []
                    position += 1
                continue
            
            # Проверка на заголовок
            if stripped.startswith('#'):
                if current_paragraph:
                    text = ' '.join(current_paragraph)
                    element = DocumentElement(
                        element_type='paragraph',
                        content=text,
                        position=position
                    )
                    elements.append(element)
                    current_paragraph = []
                    position += 1
                
                level = len(re.match(r'^#+', stripped).group())
                content_text = stripped[level:].strip()
                element = DocumentElement(
                    element_type='heading',
                    content=content_text,
                    level=min(level, 6),
                    position=position
                )
                elements.append(element)
                position += 1
            else:
                current_paragraph.append(stripped)
        
        if current_paragraph:
            text = ' '.join(current_paragraph)
            element = DocumentElement(
                element_type='paragraph',
                content=text,
                position=position
            )
            elements.append(element)
        
        return ParsedDocument(
            title=title,
            elements=elements,
            metadata={'type': 'content'}
        )
